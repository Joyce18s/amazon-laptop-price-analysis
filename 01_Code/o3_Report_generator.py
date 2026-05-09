import o2_Analysis as an
import pandas as pd
from docx2pdf import convert
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

doc = Document()
# Main Heading
MainHeading = doc.add_heading("Laptop Price Analysis", level=0)
MainHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# writing name and date
heading0 = doc.add_heading("Prepared by: Joyce Samuel", level=2)
heading1 = doc.add_heading("Date: 5/5/2026", level = 2)

# writing introduction paragraph
heading2 = doc.add_heading("Introduction", level=1)

doc.add_paragraph("This report presents an analysis of laptop price data collected " \
"from Amazon using web scraping techniques. " \
"The dataset includes various laptop models along with their prices.")
doc.add_paragraph("The objective of this report is to understand price distribution, identify the most " \
"expensive and cheapest laptops, and analyze average pricing trends across different brands.")

# writing Dataset overview
heading3 = doc.add_heading("Dataset Overview", level=1)
doc.add_paragraph(f"Total number of products: {an.tnop}"\
                 f"\nPrice range: ₹{an.min_price} to ₹{an.max_price}"\
                 f"\nAverage price: ₹{an.avg_price}")

# top 10 cheapest product
heading4 = doc.add_heading("Top 10 Cheapest Laptops", level=1)
table = doc.add_table(rows = 11, cols = 2)
table.style = "Table Grid"

table.cell(0, 0).text = "Name"
table.cell(0, 1).text = "Price"

for i, row in an.top10_cheapest.iterrows():
    table.cell(i+1, 0).text = str(row["PRODUCT NAME"])
    table.cell(i+1, 1).text = f"₹{row['PRICE']}"


# top 10 Expensive product
heading5 = doc.add_heading("Top 10 Expensive Laptops", level=1)
table = doc.add_table(rows = 11, cols = 2)
table.style = "Table Grid"

table.cell(0, 0).text = "Name"
table.cell(0, 1).text = "Price"

for i, row in an.top10_expensive.iterrows():
    table.cell(i+1, 0).text = str(row["PRODUCT NAME"])
    table.cell(i+1, 1).text = f"₹{row['PRICE']}"

# inserting boxplot graph
heading6 = doc.add_heading("Price Statistics")    
doc.add_paragraph("The boxplot shows the statics of laptop prices.")
doc.add_picture(r"C:\Users\INTEL\Desktop\Work\Amazon_Scraper\03_Charts\boxplot.png", width=Inches(7))

# inserting histogram graph
heading7 = doc.add_heading("Price Distribution")    
doc.add_paragraph("The histogram represents the comparison of laptop prices across selected products.")
doc.add_picture(r"C:\Users\INTEL\Desktop\Work\Amazon_Scraper\03_Charts\histogram.png", width=Inches(7))

# writing average price by company
heading8 = doc.add_heading("Average Price by Brand")
table = doc.add_table(rows = 13, cols = 2)
table.style = "Table Grid"

table.cell(0, 0).text = "Brand"
table.cell(0, 1).text = "Average Price"

for i, row in an.avg_price_brand.iterrows():
    table.cell(i+1, 0).text = str(row["Brand Name"])
    table.cell(i+1, 1).text = f"₹{row['Average Price']}"
doc.add_paragraph("It can be observed that Apple and Alienware laptops have the highest average price, while brands like Lenvo, Acer and Hp offer more budget-friendly options.")

# Writing key insights
heading9 = doc.add_heading("Key Insights")
doc.add_paragraph("- Most laptops are priced between ₹20,000 and ₹1,00,000."\
"- A few premium laptops exceed ₹3,00,000, acting as outliers."\
"- Lenovo and HP dominate the dataset in terms of number of products."
"- Budget laptops are mostly under ₹80,000.")
doc.add_picture(r"C:\Users\INTEL\Desktop\Work\Amazon_Scraper\03_Charts\Brand Distribution.png", width=Inches(7))

# Conclusion
heading10 = doc.add_heading("Conclusion")
doc.add_paragraph("In conclusion, the analysis shows a wide variation in laptop prices depending on brand, specifications, and category."
"This data can help customers make informed purchasing decisions and understand market trends.")

doc.save("Price Analysis Report.docx")