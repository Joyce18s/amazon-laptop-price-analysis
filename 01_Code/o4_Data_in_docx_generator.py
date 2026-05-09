from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from docx2pdf import convert

data = pd.read_excel(r"C:\Users\INTEL\Desktop\Work\Amazon_Scraper\02_Data\Laptop_Data.xlsx")
lst = []

doc = Document()
MainHeading = doc.add_heading("Laptop Price Data", level = 0)
MainHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER

table = doc.add_table(rows = len(data)+1, cols = 4)
table.style = "Table Grid"

table.cell(0, 0).text = "Index"
table.cell(0, 1).text = "Name"
table.cell(0, 2).text = "Price"
table.cell(0, 3).text = "Brand"

for i, row in data.iterrows():
    table.cell(i+1, 0).text = f"{i+1}"
    table.cell(i+1, 1).text = str(row["PRODUCT NAME"])
    table.cell(i+1, 2).text = f"Rs.{row['PRICE']}"
    table.cell(i+1, 3).text = str(row["BRAND"])
    
doc.save("Laptop_Data.docx")
convert("Laptop_Data.docx")